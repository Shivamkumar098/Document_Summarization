import vertexai
import fitz
import os
import warnings
import random
import re
import string
 
 
from vertexai.language_models import TextGenerationModel
from google.cloud import firestore
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from flask import Flask,redirect,url_for,render_template,request
app=Flask(__name__)
 
@app.route('/')
def welcome():
    return render_template('index.html')
 
 
@app.route('/login',methods=['POST','GET'])
def login():
   
    return render_template('login.html')
 
   
 
 
 
 
random.seed(42)
 
 
os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "C:\\Users\\shivam.d\\decoded-battery-389704-43834d6b2787.json"
 
 
# Initialize the Firestore client
db = firestore.Client()
 
 
# Define the common_analysis variable
common_analysis = "This is a common analysis for both documents."
 
 
# Define the reference to the Firestore document where you want to store the data
doc_ref = db.collection('bison').document('oHBvRPOIvGrv5iFlbCBF')
 
 
# Initialize a list to store document results
document_results = []
 
 
# Function for combined preprocessing steps
def preprocess_text(text):
 
    # Remove special characters and extra whitespace
    text = re.sub(r'[^\w\s]', '', text)
 
    text = re.sub(r'\s+', ' ', text)
 
    # Convert to lowercase
    text = text.lower()
 
    # Remove punctuation
    text = text.translate(str.maketrans('', '', string.punctuation))
 
    # Remove emojis and other non-ASCII characters
    text = text.encode("ascii", "ignore").decode()
 
    # Tokenization
    tokens = word_tokenize(text)
 
    # Remove stopwords
    stop_words = set(stopwords.words('english'))
 
    filtered_tokens = [word for word in tokens if word not in stop_words]
 
    # Lemmatization
    lemmatizer = WordNetLemmatizer()
 
    lemmatized_tokens = [lemmatizer.lemmatize(word) for word in filtered_tokens]
 
    # Remove short words
    min_word_length = 2
 
    preprocessed_text = ' '.join([word for word in lemmatized_tokens if len(word) >= min_word_length])
 
 
    return preprocessed_text
 
 
 
# Get user input for document paths
print("Enter the paths of the documents. Press Enter after each path. Enter 'done' when finished.")
 
document_paths = []
 
while True:
 
    document_path = input("Document path: ")
 
    if document_path.lower() == 'done':
        break
 
    document_paths.append(document_path)
 
 
 
model = TextGenerationModel.from_pretrained("text-bison@001")
 
 
 
# Function for generating UI responses
def get_tsg_functional_response():
 
    # Get the system name from the user
    system_name = input("Enter the system name: ")
 
    vertexai.init(project="decoded-battery-389704", location="us-central1")
 
    parameters_tsg = {
 
        "temperature": 1,
        "max_output_tokens": 300,
        "top_p": 1,
        "top_k": 40
 
    }
 
 
    # Set the prefix
    suffix = """Statement:\n\"{system_name}\nwhat all things are mandatory to be part of business story from this use case in User interface (UI) test cases: (give answer in 5 points in new line)."""
 
    # Combine the prefix and document content as the context for generating a response
    context_tsg = preprocessed_content + "\n\n" + suffix
 
    # Send the user input to the chat model and generate a response
    response_tsg = model.predict(context_tsg, **parameters_tsg)
 
 
 
    suffix_2 = """Statement:\n\"{preprocessed_content}\"\n{system_name}\nwhat all things are mandatory to be part of business story from this use case in Usability test cases: (give answer in 5 points in new line)."""
 
    usability_response = model.predict(suffix_2, **parameters_tsg)
 
 
 
    suffix_3 = """Statement:\n\"{preprocessed_content}\"\n{system_name}\nwhat all things are mandatory to be part of business story from this use case in Functionality test cases: (give answer in 5 points in new line)."""
 
    functionality_response = model.predict(suffix_3, **parameters_tsg)
 
 
 
    suffix_4 = """Statement:\n\"{preprocessed_content}\"\n{system_name}\nwhat all things are mandatory to be part of business story from this use case in Integration test cases: (give answer in 5 points in new line)."""
 
    integration_response = model.predict(suffix_4, **parameters_tsg)
 
 
 
    suffix_5 = """Statement:\n\"{preprocessed_content}\"\n{system_name}\nwhat all things are mandatory to be part of business story from this use case in Security test cases: (give answer in 5 points in new line)."""
    security_response = model.predict(suffix_5, **parameters_tsg)
 
 
 
    suffix_6 = """Statement:\n\"{preprocessed_content}\"\n{system_name}\nwhat all things are mandatory to be part of business story from this use case in Performance test cases: (give answer in 5 points in new line)."""
 
    performance_response = model.predict(suffix_6, **parameters_tsg)
 
 
 
    ui_response = response_tsg.text
 
 
    usability_response = usability_response.text
 
 
    functionality_response = functionality_response.text
 
 
    integration_response = integration_response.text
 
 
    security_response = security_response.text
 
    security_response = security_response
 
 
    performance_response = performance_response.text
 
    performance_response = performance_response
 
 
 
    # Remove asterisks from UI Test Cases
    ui_response = ui_response.replace('*', '')
 
 
    # Remove asterisks from Usability Test Cases
    usability_response = usability_response.replace('*', '')
 
 
    # Remove asterisks from Functionality Test Cases
    functionality_response = functionality_response.replace('*', '')
 
 
    # Remove asterisks from Integration Test Cases
    integration_response = integration_response.replace('*', '')
 
 
    # Remove asterisks from Security Test Cases
    security_response = security_response.replace('*', '')
 
 
    # Remove asterisks from Performance Test Cases
    performance_response = performance_response.replace('*', '')
 
 
    return system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response
   
 
 
# Create lists to store test scenarios for each document
document1_test_scenarios = []
document2_test_scenarios = []
 
 
# Create variables to store document content
doc_content_1 = ""
doc_content_2 = ""
 
 
# Loop through each document
for i, document_path in enumerate(document_paths):
 
 
    # Initialize an empty dictionary to store test scenarios for the current document
    document_test_scenarios = {}
 
 
    # Initialize an empty dictionary to store page contents
    document_pages = {}
 
 
    # Initialize an empty dictionary to store results for each document
    document_result = {}
 
 
    # Open the PDF document using PyMuPDF
    doc = fitz.open(document_path)
 
 
    # Extract text content from all pages
    document_content = ""
 
    for page_num in range(doc.page_count):
 
        page = doc.load_page(page_num)
 
        text = page.get_text("text")
 
        document_content += f"Page {page_num + 1}:\n{text}\n\n"  # Adding page number and line breaks
 
 
        # Store page text in the dictionary
        document_pages[page_num] = text
 
 
    # Preprocess document content and apply additional preprocessing steps
    preprocessed_content = preprocess_text(document_content)
 
    # Store the document content in the appropriate variable
    if i == 0:
        doc_content_1 = preprocessed_content
    elif i == 1:
        doc_content_2 = preprocessed_content
 
 
    # Call the get_tsg_functional_response function and store the responses
    system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response = get_tsg_functional_response()
 
 
    # Store the test scenarios for document 1 in the first iteration and document 2 in the second iteration
    if len(document1_test_scenarios) == len(document2_test_scenarios):
 
        document1_test_scenarios.append((system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response))
   
    else:
 
        document2_test_scenarios.append((system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response))
 
 
 
# Print the test scenarios as bullet points
print("\nTest Scenarios for Document 1:")
 
for i, scenario in enumerate(document1_test_scenarios, start=1):
 
    system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response = scenario
   
    print(f"{i}. System Name: {system_name}")
 
    print("   UI Test Cases:")
    for j, ui_test_case in enumerate(ui_response.split('\n'), start=1):
        print(f"   {j}. {ui_test_case.strip('*')}")
 
    print("   Usability Test Cases:")
    for j, usability_test_case in enumerate(usability_response.split('\n'), start=1):
        print(f"   {j}. {usability_test_case.strip('*')}")
 
    print("   Functionality Test Cases:")
    for j, functionality_test_case in enumerate(functionality_response.split('\n'), start=1):
        print(f"   {j}. {functionality_test_case.strip('*')}")
 
    print("   Integration Test Cases:")
    for j, integration_test_case in enumerate(integration_response.split('\n'), start=1):
        print(f"   {j}. {integration_test_case.strip('*')}")
 
    print("   Security Test Cases:")
    for j, security_test_case in enumerate(security_response.split('\n'), start=1):
        print(f"   {j}. {security_test_case.strip('*')}")
 
    print("   Performance Test Cases:")
    for j, performance_test_case in enumerate(performance_response.split('\n'), start=1):
        print(f"   {j}. {performance_test_case.strip('*')}")
 
    print("")  
 
 
 
print("\nTest Scenarios for Document 2:")
 
for i, scenario in enumerate(document2_test_scenarios, start=1):
 
    system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response = scenario
   
    print(f"{i}. System Name: {system_name}")
 
    print("   UI Test Cases:")
    for j, ui_test_case in enumerate(ui_response.split('\n'), start=1):
        print(f"   {j}. {ui_test_case.strip('*')}")
 
    print("   Usability Test Cases:")
    for j, usability_test_case in enumerate(usability_response.split('\n'), start=1):
        print(f"   {j}. {usability_test_case.strip('*')}")
 
    print("   Functionality Test Cases:")
    for j, functionality_test_case in enumerate(functionality_response.split('\n'), start=1):
        print(f"   {j}. {functionality_test_case.strip('*')}")
 
    print("   Integration Test Cases:")
    for j, integration_test_case in enumerate(integration_response.split('\n'), start=1):
        print(f"   {j}. {integration_test_case.strip('*')}")
 
    print("   Security Test Cases:")
    for j, security_test_case in enumerate(security_response.split('\n'), start=1):
        print(f"   {j}. {security_test_case.strip('*')}")
 
    print("   Performance Test Cases:")
    for j, performance_test_case in enumerate(performance_response.split('\n'), start=1):
        print(f"   {j}. {performance_test_case.strip('*')}")
 
    print("")  # Empty line to separate scenarios
 
 
 
vertexai.init(project="decoded-battery-389704", location="us-central1")
 
 
parameters = {
 
    "temperature": 0.2,
    "max_output_tokens": 256,
    "top_p": 0.8,
    "top_k": 40
 
    }
 
 
# Set the prefix
context_diff = """Statement:\n\"{document1_test_scenarios}\"\n and \n{document2_test_scenarios}\nAnalyze these two test scenarios which are for two different design documents and find and list out the points which are different in the test scenarios for both documents under the heading of Regression. And give answer in more technical terms with specific and detailed information. (give answer in 5  points in new line)."""
 
 
# Send the user input to the chat model and generate a response
response_diff = model.predict(context_diff, **parameters)
 
 
analysis = response_diff.text
 
 
# After generating the analysis content, remove asterisks from it as well
analysis = analysis.replace('*', '')
 
 
# Print the generated response
print("Regression ")
 
print(analysis)
 
 
# Close the PDF document
doc.close()
 
 
doc = Document()
 
 
# Function to remove empty lines from a text block
def remove_empty_lines(text):
 
    # Split the text into lines, filter out empty lines, and join them with line breaks
    cleaned_lines = [line for line in text.splitlines() if line.strip()]
 
    cleaned_text = '\n'.join(cleaned_lines)
 
    return cleaned_text
 
 
def add_bullet_points(section_title, test_scenarios):
 
    # Add a section heading
    section_heading = doc.add_heading(section_title, level=2)
 
    section_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
 
 
    # Add each test scenario as a bullet point without "**"
    for i, scenario in enumerate(test_scenarios, start=1):
 
        system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response = scenario
 
 
        # Add the system name as a bullet point
        doc.add_paragraph(f"{i}. System Name: {system_name}", style='List Bullet')
 
 
        # Add UI Test Cases
        ui_response_lines = [line for line in ui_response.split('\n') if line.strip()]
 
        doc.add_paragraph('UI Test Cases:', style='List Bullet')
 
        for ui_test_case in ui_response_lines:
 
            doc.add_paragraph(ui_test_case, style='List Bullet')
 
 
        # Add Usability Test Cases
        usability_response_lines = [line for line in usability_response.split('\n') if line.strip()]
 
        doc.add_paragraph('Usability Test Cases:', style='List Bullet')
 
        for usability_test_case in usability_response_lines:
 
            doc.add_paragraph(usability_test_case, style='List Bullet')
 
 
        # Add Functionality Test Cases
        functionality_response_lines = [line for line in functionality_response.split('\n') if line.strip()]
 
        doc.add_paragraph('Functionality Test Cases:', style='List Bullet')
 
        for functionality_test_case in functionality_response_lines:
 
            doc.add_paragraph(functionality_test_case, style='List Bullet')
 
 
        # Add Integration Test Cases
        integration_response_lines = [line for line in integration_response.split('\n') if line.strip()]
 
        doc.add_paragraph('Integration Test Cases:', style='List Bullet')
 
        for integration_test_case in integration_response_lines:
 
            doc.add_paragraph(integration_test_case, style='List Bullet')
 
 
        # Add Security Test Cases
        security_response_lines = [line for line in security_response.split('\n') if line.strip()]
 
        doc.add_paragraph('Security Test Cases:', style='List Bullet')
 
        for security_test_case in security_response_lines:
 
            doc.add_paragraph(security_test_case, style='List Bullet')
 
 
        # Add Performance Test Cases
        performance_response_lines = [line for line in performance_response.split('\n') if line.strip()]
 
        doc.add_paragraph('Performance Test Cases:', style='List Bullet')
 
        for performance_test_case in performance_response_lines:
 
            doc.add_paragraph(performance_test_case, style='List Bullet')
 
 
        # Add an empty line to separate scenarios
        doc.add_paragraph('')
 
 
 
# When adding test scenarios for Document 1, modify the section title as follows:
add_bullet_points(f"Test Scenarios for :- ", document1_test_scenarios)
 
 
# When adding test scenarios for Document 2, modify the section title as follows:
add_bullet_points(f"Test Scenarios for :-", document2_test_scenarios)
 
 
 
analysis_heading = doc.add_heading('Regression', level=1)
 
analysis_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
 
 
 
# Add the analysis content (you may want to format it as needed)
analysis_paragraph = doc.add_paragraph(analysis)
 
analysis_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
 
 
 
# Save the Word document to a specific directory (provide the full path)
output_path = 'C:\\Users\\shivam.d\\Documents\\pdf\\Test_scenarios_and_analysis44.docx'
 
doc.save(output_path)
 
 
print(f"Word document saved to '{output_path}'")
 
 
 
# Define the Firestore collection and document ID
collection_name = 'Test_scenarios'
 
document_id = 'YH6TZjJDDCzq5mSJnYTG'
 
 
# Convert document1_test_scenarios and document2_test_scenarios lists to dictionaries
document1_data = {
 
    f'System Name: {system_name}': {
 
        'UI Test Cases': ui_response.split('\n'),
 
        'Usability Test Cases': usability_response.split('\n'),
 
        'Functionality Test Cases': functionality_response.split('\n'),
 
        'Integration Test Cases': integration_response.split('\n'),
 
        'Security Test Cases': security_response.split('\n'),
 
        'Performance Test Cases': performance_response.split('\n')
 
    }
 
    for system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response in document1_test_scenarios
 
}
 
 
document2_data = {
 
    f'System Name: {system_name}': {
 
        'UI Test Cases': ui_response.split('\n'),
 
        'Usability Test Cases': usability_response.split('\n'),
 
        'Functionality Test Cases': functionality_response.split('\n'),
 
        'Integration Test Cases': integration_response.split('\n'),
 
        'Security Test Cases': security_response.split('\n'),
 
        'Performance Test Cases': performance_response.split('\n')
 
    }
 
    for system_name, ui_response, usability_response, functionality_response, integration_response, security_response, performance_response in document2_test_scenarios
 
}
 
 
 
for system_name, test_cases in document1_data.items():
 
    for test_case_type, test_case_list in test_cases.items():
 
        # Remove empty points from each test case list
        cleaned_test_case_list = [test_case.strip() for test_case in test_case_list if test_case.strip()]
 
        # Update the test case list in the dictionary
        document1_data[system_name][test_case_type] = cleaned_test_case_list
 
 
for system_name, test_cases in document2_data.items():
 
    for test_case_type, test_case_list in test_cases.items():
 
        # Remove empty points from each test case list
        cleaned_test_case_list = [test_case.strip() for test_case in test_case_list if test_case.strip()]
 
        # Update the test case list in the dictionary
        document2_data[system_name][test_case_type] = cleaned_test_case_list
 
 
# Define the data you want to store in Firestore
data = {
 
    'Test_Scenarios (Doc1)': document1_data,
 
    'Test_Scenarios (Doc2)': document2_data,
 
    'Doc 1_Content': doc_content_1,
 
    'Doc 2_Content': doc_content_2,
 
    'Regression': analysis,
 
}
 
 
db = firestore.Client()
 
 
# Define the Firestore document reference where you want to store the data
doc_ref = db.collection(collection_name).document(document_id)
 
 
try:
 
    # Set the combined data in Firestore
    doc_ref.set(data)
 
    print("Test Analysis and Regression stored in Firestore with ID:", doc_ref.id)
 
except Exception as e:
    print("Error while storing data in Firestore:", str(e))
 
 
 
# Filter and suppress specific warnings
warnings.filterwarnings("ignore", message="Exception ignored in: <function Outline.__del__.*")
 
 
# Temporarily suppress warnings from the specified module
with warnings.catch_warnings():
 
    warnings.simplefilter("ignore", category=UserWarning)  # You might need to adjust the warning category
 
 
 
 
if __name__=='__main__':
    app.run(port=8080)
   
# # Function to save data to Firestore
# def save_to_firestore(user_input, test_scenarios, analysis):
#     # Define the Firestore collection and document ID
#     collection_name = 'Test_scenarios'
#     document_id = 'YH6TZjJDDCzq5mSJnYTG'  # Use your desired document ID
 
#     # Get the existing data from Firestore
#     doc_ref = db.collection(collection_name).document(document_id)
#     existing_data = doc_ref.get().to_dict()
 
#     # Create or update a new field with user-generated data
#     new_field_name = f'user_input_{len(existing_data) + 1}'  # Create a new field name
#     data = {
#         'Test_Scenarios (Doc1)': document1_data,
#         'Test_Scenarios (Doc2)': document2_data,
#         'regression': analysis
#     }
#     existing_data[new_field_name] = new_data
 
#     # Update the Firestore document with the new data
#     doc_ref.update(existing_data)
 
#     print("Data saved to Firestore with ID:", doc_ref.id)
 
 
 
# # Get user input for document paths
 
 
# while True:
#     user_input = input("Enter your input: ")
#     if user_input.lower() == 'done':
#         break
 
 
#     # Save the data to Firestore
#     save_to_firestore(user_input, document1_test_scenarios, analysis)